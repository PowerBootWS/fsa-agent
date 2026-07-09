"""
Resume/cover-letter tailoring agent — generates a job-specific tailored resume and/or
cover letter from a candidate's uploaded resume (and optional cover letter) plus a saved
job's snapshot. Uses OpenRouter directly (not the claude -p pattern used elsewhere in this
stack) so usage/cost for this customer-facing, potentially high-volume feature stays
isolated from the Claude Code Max subscription's own limits — an owner-approved, deliberate
exception (2026-07-08).

Unlike checkpoint.py's fallback-on-error pattern (fine for a throwaway chat nudge),
generation failures here propagate to the caller — this produces a paid deliverable the
candidate downloads, so silently returning placeholder content on an API error would be
worse than a visible failure.
"""
import json
import os
import re

import pdfplumber
import requests
from docx import Document as DocxDocument
from weasyprint import HTML

_OPENROUTER_URL = 'https://openrouter.ai/api/v1/chat/completions'

_SYSTEM = (
    "You are an expert resume and cover letter writer for Power Engineering job "
    "candidates. You tailor a candidate's existing resume/cover letter to a specific "
    "job posting. Follow these rules strictly:\n"
    "1. Never invent facts, dates, employers, certifications, or numbers that are not "
    "already present in the candidate's source material. If a strong bullet point needs "
    "information the candidate hasn't provided, insert a bracketed placeholder instead "
    "(e.g. '[specific boiler type/manufacturer]', '[reason you want to work here]') — "
    "target 3 to 6 placeholders total across the requested documents.\n"
    "2. Mirror the job posting's own terminology and reorder/re-emphasize the candidate's "
    "existing experience to match what the posting stresses — do not add experience that "
    "isn't there.\n"
    "3. If the posting has a hard requirement the candidate's background doesn't clearly "
    "meet (e.g. a certification level), note it in changes_summary or flagged_gaps — "
    "never paper over it in the document itself.\n"
    "4. changes_summary must be first-person, plain-language coaching directly to the "
    "candidate (e.g. \"I moved your certificate to the top...\") — never phrased as "
    "compliance against a reference document.\n"
    "Respond with ONLY a JSON object matching this shape, no other text:\n"
    '{"resume_content": "...or null", "cover_letter_content": "...or null", '
    '"changes_summary": "...", "placeholder_count": 0, "flagged_gaps": ["..."]}'
)


def extract_text(file_path: str, mime_type: str) -> str:
    """Extract plain text from an uploaded resume/cover-letter file (PDF or DOCX)."""
    if mime_type == 'application/pdf':
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        text = '\n'.join(text_parts).strip()
    else:
        doc = DocxDocument(file_path)
        text = '\n'.join(p.text for p in doc.paragraphs).strip()

    if not text:
        raise ValueError('NO_EXTRACTABLE_TEXT')
    return text


def _build_prompt(job: dict, resume_text: str, cover_letter_text: str | None, doc_types: list[str]) -> str:
    wants = ' and '.join(
        {'resume': 'a tailored resume', 'cover_letter': 'a tailored cover letter'}[t] for t in doc_types
    )
    parts = [
        f"Generate {wants} for this candidate, tailored to the job posting below.",
        f"\nJob title: {job.get('title', '')}",
        f"Company: {job.get('company', '')}",
        f"Job posting description:\n{job.get('description_snapshot', '')}",
        f"\nCandidate's current resume:\n{resume_text}",
    ]
    if cover_letter_text:
        parts.append(f"\nCandidate's current cover letter (for tone/history context only):\n{cover_letter_text}")
    if 'resume' not in doc_types:
        parts.append('\nSet resume_content to null — a tailored resume was not requested.')
    if 'cover_letter' not in doc_types:
        parts.append('\nSet cover_letter_content to null — a tailored cover letter was not requested.')
    return '\n'.join(parts)


def _call_openrouter(prompt: str) -> tuple[dict, str]:
    api_key = os.getenv('OPENROUTER_API_KEY')
    # Confirm this exact slug against OpenRouter's current model list before deploy
    # (spec open question #1, docs/superpowers/specs/2026-07-08-resume-cover-letter-tailoring-design.md).
    model = os.getenv('OPENROUTER_RESUME_MODEL', 'anthropic/claude-sonnet-5')

    response = requests.post(
        _OPENROUTER_URL,
        headers={
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json',
        },
        json={
            'model': model,
            'max_tokens': 4000,
            'messages': [
                {'role': 'system', 'content': _SYSTEM},
                {'role': 'user', 'content': prompt},
            ],
        },
        timeout=60,
    )
    response.raise_for_status()
    raw_content = response.json()['choices'][0]['message']['content'].strip()

    # Some models wrap JSON in ```json fences despite instructions — strip if present.
    match = re.search(r'\{.*\}', raw_content, re.DOTALL)
    if not match:
        raise ValueError('MODEL_RESPONSE_NOT_JSON')
    parsed = json.loads(match.group(0))
    return parsed, model


def _render_docx(title: str, content: str, output_path: str) -> None:
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    for line in content.split('\n'):
        if not line.strip():
            continue
        paragraph = doc.add_paragraph()
        # Bracketed placeholders are bolded so they stand out for the candidate to find
        # and fill in before sending the document anywhere.
        segments = re.split(r'(\[[^\]]+\])', line)
        for segment in segments:
            run = paragraph.add_run(segment)
            if segment.startswith('[') and segment.endswith(']'):
                run.bold = True
    doc.save(output_path)


def _render_pdf(title: str, content: str, output_path: str) -> None:
    def _escape(text):
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    escaped_lines = [
        re.sub(r'\[[^\]]+\]', lambda m: f'<strong>{m.group(0)}</strong>', _escape(line))
        for line in content.split('\n')
        if line.strip()
    ]
    html = f"""
    <html><head><style>
      body {{ font-family: sans-serif; font-size: 11pt; line-height: 1.5; margin: 2cm; }}
      h1 {{ font-size: 16pt; }}
      strong {{ background: #fff3b0; }}
    </style></head>
    <body><h1>{_escape(title)}</h1>{''.join(f'<p>{line}</p>' for line in escaped_lines)}</body>
    </html>
    """
    HTML(string=html).write_pdf(output_path)


def generate_tailored_documents(job: dict, resume_path: str, resume_mime: str,
                                 cover_letter_path: str | None, cover_letter_mime: str | None,
                                 doc_types: list[str], output_dir: str) -> dict:
    """
    Orchestrates extraction -> OpenRouter call -> rendering for one tailoring request.
    Raises ValueError('NO_EXTRACTABLE_TEXT') if the source resume/cover letter can't be
    read, or ValueError('MODEL_DID_NOT_RETURN_<TYPE>') if the model didn't produce a
    requested document type.
    """
    resume_text = extract_text(resume_path, resume_mime)
    cover_letter_text = extract_text(cover_letter_path, cover_letter_mime) if cover_letter_path else None

    prompt = _build_prompt(job, resume_text, cover_letter_text, doc_types)
    parsed, model_used = _call_openrouter(prompt)

    os.makedirs(output_dir, exist_ok=True)
    content_by_type = {
        'resume': (f"Resume — {job.get('title', '')}", parsed.get('resume_content')),
        'cover_letter': (f"Cover Letter — {job.get('title', '')}", parsed.get('cover_letter_content')),
    }
    documents = []
    for doc_type in doc_types:
        title, content = content_by_type[doc_type]
        if not content:
            raise ValueError(f'MODEL_DID_NOT_RETURN_{doc_type.upper()}')
        docx_path = os.path.join(output_dir, f'{doc_type}.docx')
        pdf_path = os.path.join(output_dir, f'{doc_type}.pdf')
        _render_docx(title, content, docx_path)
        _render_pdf(title, content, pdf_path)
        documents.append({'doc_type': doc_type, 'docx_path': docx_path, 'pdf_path': pdf_path})

    return {
        'documents': documents,
        'changes_summary': parsed.get('changes_summary', ''),
        'placeholder_count': int(parsed.get('placeholder_count', 0)),
        'flagged_gaps': parsed.get('flagged_gaps', []),
        'model_used': model_used,
    }
