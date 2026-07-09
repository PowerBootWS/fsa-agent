import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from agents.resume_writer import extract_text

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), 'fixtures')


def test_extract_text_from_docx():
    text = extract_text(os.path.join(FIXTURES_DIR, 'sample_resume.docx'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    assert 'Jane Candidate' in text
    assert 'package boilers' in text


def test_extract_text_from_pdf():
    text = extract_text(os.path.join(FIXTURES_DIR, 'sample_resume.pdf'), 'application/pdf')
    assert 'Jane Candidate' in text


def test_extract_text_raises_on_empty_document():
    import tempfile
    from docx import Document
    doc = Document()
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc.save(f.name)
        path = f.name
    try:
        import pytest
        with pytest.raises(ValueError, match='NO_EXTRACTABLE_TEXT'):
            extract_text(path, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    finally:
        os.unlink(path)


from unittest.mock import patch, MagicMock
from agents.resume_writer import _build_prompt, _call_openrouter


def test_build_prompt_includes_job_and_resume_and_omits_unrequested_types():
    job = {'title': 'Boiler Operator', 'company': 'Acme Plant', 'description_snapshot': 'Operate package boilers.'}
    prompt = _build_prompt(job, 'Jane Candidate, 3rd Class.', None, ['resume'])
    assert 'Boiler Operator' in prompt
    assert 'Jane Candidate, 3rd Class.' in prompt
    assert 'cover_letter_content to null' in prompt


def test_build_prompt_includes_cover_letter_context_when_provided():
    job = {'title': 'Boiler Operator', 'company': 'Acme Plant', 'description_snapshot': 'Operate package boilers.'}
    prompt = _build_prompt(job, 'Jane Candidate.', 'Dear Hiring Manager, existing cover letter text.', ['resume', 'cover_letter'])
    assert 'existing cover letter text' in prompt


@patch('agents.resume_writer.requests.post')
def test_call_openrouter_parses_json_response(mock_post):
    mock_response = MagicMock()
    mock_response.json.return_value = {
        'choices': [{'message': {'content': '{"resume_content": "Tailored resume text", "cover_letter_content": null, "changes_summary": "Moved cert to top.", "placeholder_count": 4, "flagged_gaps": []}'}}]
    }
    mock_response.raise_for_status.return_value = None
    mock_post.return_value = mock_response

    parsed, model_used = _call_openrouter('some prompt')

    assert parsed['resume_content'] == 'Tailored resume text'
    assert parsed['placeholder_count'] == 4
    assert model_used  # non-empty model id was used


@patch('agents.resume_writer.requests.post')
def test_call_openrouter_strips_markdown_fences(mock_post):
    mock_response = MagicMock()
    mock_response.json.return_value = {
        'choices': [{'message': {'content': '```json\n{"resume_content": "x", "cover_letter_content": null, "changes_summary": "y", "placeholder_count": 1, "flagged_gaps": []}\n```'}}]
    }
    mock_response.raise_for_status.return_value = None
    mock_post.return_value = mock_response

    parsed, _ = _call_openrouter('some prompt')
    assert parsed['resume_content'] == 'x'


@patch('agents.resume_writer.requests.post')
def test_call_openrouter_raises_on_non_json_response(mock_post):
    import pytest
    mock_response = MagicMock()
    mock_response.json.return_value = {'choices': [{'message': {'content': 'Sorry, I cannot help with that.'}}]}
    mock_response.raise_for_status.return_value = None
    mock_post.return_value = mock_response

    with pytest.raises(ValueError, match='MODEL_RESPONSE_NOT_JSON'):
        _call_openrouter('some prompt')


import tempfile
import shutil
from docx import Document as DocxDocument
from agents.resume_writer import _render_docx, _render_pdf, generate_tailored_documents


def test_render_docx_bolds_bracketed_placeholders():
    tmp_dir = tempfile.mkdtemp()
    try:
        output_path = os.path.join(tmp_dir, 'resume.docx')
        _render_docx('Resume — Boiler Operator', 'Experience with [specific boiler type/manufacturer].', output_path)
        assert os.path.exists(output_path)
        doc = DocxDocument(output_path)
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        assert '[specific boiler type/manufacturer]' in full_text
    finally:
        shutil.rmtree(tmp_dir)


def test_render_pdf_creates_a_readable_file():
    tmp_dir = tempfile.mkdtemp()
    try:
        output_path = os.path.join(tmp_dir, 'resume.pdf')
        _render_pdf('Resume — Boiler Operator', 'Experience with [specific boiler type/manufacturer].', output_path)
        assert os.path.exists(output_path)
        assert os.path.getsize(output_path) > 0
    finally:
        shutil.rmtree(tmp_dir)


@patch('agents.resume_writer._call_openrouter')
def test_generate_tailored_documents_returns_paths_and_summary(mock_call):
    mock_call.return_value = (
        {
            'resume_content': 'Tailored resume body with a [placeholder].',
            'cover_letter_content': None,
            'changes_summary': 'Moved certification to the top.',
            'placeholder_count': 1,
            'flagged_gaps': [],
        },
        'anthropic/claude-sonnet-5',
    )
    tmp_dir = tempfile.mkdtemp()
    try:
        job = {'title': 'Boiler Operator', 'company': 'Acme Plant', 'description_snapshot': 'Operate boilers.'}
        result = generate_tailored_documents(
            job=job,
            resume_path=os.path.join(FIXTURES_DIR, 'sample_resume.docx'),
            resume_mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            cover_letter_path=None,
            cover_letter_mime=None,
            doc_types=['resume'],
            output_dir=tmp_dir,
        )
        assert result['changes_summary'] == 'Moved certification to the top.'
        assert result['placeholder_count'] == 1
        assert result['model_used'] == 'anthropic/claude-sonnet-5'
        assert len(result['documents']) == 1
        doc = result['documents'][0]
        assert doc['doc_type'] == 'resume'
        assert os.path.exists(doc['docx_path'])
        assert os.path.exists(doc['pdf_path'])
    finally:
        shutil.rmtree(tmp_dir)


@patch('agents.resume_writer._call_openrouter')
def test_generate_tailored_documents_raises_if_model_omits_a_requested_type(mock_call):
    import pytest
    mock_call.return_value = (
        {'resume_content': None, 'cover_letter_content': None, 'changes_summary': '', 'placeholder_count': 0, 'flagged_gaps': []},
        'anthropic/claude-sonnet-5',
    )
    tmp_dir = tempfile.mkdtemp()
    try:
        job = {'title': 'Boiler Operator', 'company': 'Acme Plant', 'description_snapshot': 'Operate boilers.'}
        with pytest.raises(ValueError, match='MODEL_DID_NOT_RETURN_RESUME'):
            generate_tailored_documents(
                job=job,
                resume_path=os.path.join(FIXTURES_DIR, 'sample_resume.docx'),
                resume_mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                cover_letter_path=None,
                cover_letter_mime=None,
                doc_types=['resume'],
                output_dir=tmp_dir,
            )
    finally:
        shutil.rmtree(tmp_dir)
